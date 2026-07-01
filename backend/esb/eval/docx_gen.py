"""Generate ESB Board Meeting Time Use Evaluation Word document.

Matches the ESB template exactly:
  - Margins: 0.5" all sides (720 twips)
  - Header: centered ESB logo (2.67" × 0.68")
  - Footer: centered gray italic website + tagline, with page numbers
"""
from pathlib import Path
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

ESB_DARK = RGBColor(0x1F, 0x4E, 0x79)
ESB_BLUE = RGBColor(0x2E, 0x75, 0xB6)
ESB_LIGHT = RGBColor(0xD5, 0xE8, 0xF0)
ESB_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ESB_GRAY = RGBColor(0x64, 0x64, 0x64)

# Path to the logo bundled with the app
LOGO_PATH = Path(__file__).parent / "assets" / "esb_logo.png"


def _set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_page_number_field(run):
    """Insert a PAGE field into a run."""
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    run._r.append(instr_text)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


def _add_num_pages_field(run):
    """Insert a NUMPAGES field into a run."""
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " NUMPAGES "
    run._r.append(instr_text)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


def _setup_header(section):
    """Add centered ESB logo to header, matching the template."""
    header = section.header
    # Clear default empty paragraph
    for para in header.paragraphs:
        p = para._p
        p.getparent().remove(p)

    para = header.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = para._p.get_or_add_pPr()
    # Match template spacing: after=80, line=276
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "276")
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)

    if LOGO_PATH.exists():
        run = para.add_run()
        # Image dimensions from template: cx=2560320, cy=650573 EMU
        run.add_picture(str(LOGO_PATH), width=Emu(2560320), height=Emu(650573))
    else:
        # Fallback text if logo missing
        run = para.add_run("Effective School Boards")
        run.bold = True
        run.font.color.rgb = ESB_DARK
        run.font.size = Pt(14)


def _setup_footer(section):
    """Add footer matching template + page numbers.

    Template text: www.effectiveschoolboards.com  •  Student outcomes don't change until adult behaviors change.™
    Addition: page number (e.g. 'Page 1 of 3')
    """
    footer = section.footer
    for para in footer.paragraphs:
        p = para._p
        p.getparent().remove(p)

    # Line 1: template text
    para1 = footer.add_paragraph()
    para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = para1._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "288")
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)

    def _gray_italic_run(para, text):
        r = para.add_run(text)
        r.italic = True
        r.font.color.rgb = ESB_GRAY
        r.font.size = Pt(8)
        return r

    def _gray_run(para, text):
        r = para.add_run(text)
        r.font.color.rgb = ESB_GRAY
        r.font.size = Pt(8)
        return r

    _gray_italic_run(para1, "www.effectiveschoolboards.com")
    _gray_run(para1, "   •   ")
    _gray_italic_run(para1, "Student outcomes don’t change until adult behaviors change.™")

    # Line 2: page number
    para2 = footer.add_paragraph()
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr2 = para2._p.get_or_add_pPr()
    spacing2 = OxmlElement("w:spacing")
    spacing2.set(qn("w:before"), "0")
    spacing2.set(qn("w:line"), "240")
    spacing2.set(qn("w:lineRule"), "auto")
    pPr2.append(spacing2)

    page_run = para2.add_run("Page ")
    page_run.font.color.rgb = ESB_GRAY
    page_run.font.size = Pt(7)

    page_num_run = para2.add_run()
    page_num_run.font.color.rgb = ESB_GRAY
    page_num_run.font.size = Pt(7)
    _add_page_number_field(page_num_run)

    of_run = para2.add_run(" of ")
    of_run.font.color.rgb = ESB_GRAY
    of_run.font.size = Pt(7)

    total_run = para2.add_run()
    total_run.font.color.rgb = ESB_GRAY
    total_run.font.size = Pt(7)
    _add_num_pages_field(total_run)


def _section_heading(doc: Document, text: str):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = ESB_DARK
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(3)


def _header_row(table, texts: List[str]):
    row = table.rows[0]
    for cell, text in zip(row.cells, texts):
        _set_cell_bg(cell, ESB_DARK)
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(text)
        run.bold = True
        run.font.color.rgb = ESB_WHITE
        run.font.size = Pt(9)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_info_row(table, label: str, value: str):
    row = table.add_row()
    label_cell, value_cell = row.cells[0], row.cells[1]
    _set_cell_bg(label_cell, ESB_LIGHT)
    lp = label_cell.paragraphs[0]
    lp.clear()
    lr = lp.add_run(label)
    lr.bold = True
    lr.font.size = Pt(9)
    lr.font.color.rgb = ESB_DARK

    vp = value_cell.paragraphs[0]
    vp.clear()
    vr = vp.add_run(value or "")
    vr.font.size = Pt(9)


def generate_evaluation_docx(
    output_path: str,
    district: str,
    meeting_date: str,
    meeting_type: str,
    video_url: str,
    agenda_link: str,
    strategic_plan_link: str,
    agenda_items: List[Dict],
    classified_items: List[Dict],
    category_totals: Dict[str, Any],
    coaching: Dict[str, List[Dict]],
):
    doc = Document()

    # ── Margins: 0.5" all sides, matching template (720 twips = 0.5")  ──────
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.3)
        _setup_header(section)
        _setup_footer(section)

    # ── Title ──────────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(2)
    tr = title.add_run("ESB Board Meeting Time Use Evaluation")
    tr.bold = True
    tr.font.size = Pt(14)
    tr.font.color.rgb = ESB_DARK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(6)
    sr = subtitle.add_run(district)
    sr.font.size = Pt(11)
    sr.font.color.rgb = ESB_BLUE

    # ── Section 1: Meeting Information ────────────────────────────────────────
    _section_heading(doc, "Section 1: Meeting Information")
    t1 = doc.add_table(rows=1, cols=2)
    t1.style = "Table Grid"
    t1.columns[0].width = Inches(2.2)
    t1.columns[1].width = Inches(4.8)
    _header_row(t1, ["Field", "Value"])

    total_min = category_totals["total_minutes"]
    gf_min = category_totals["goal_focused_minutes"]
    gf_pct = category_totals["goal_focused_pct"]

    rows_data = [
        ("School System", district),
        ("Meeting Date", meeting_date),
        ("Meeting Type", meeting_type),
        ("Video / Transcript Link", video_url),
        ("Agenda Link", agenda_link or ""),
        ("Strategic Plan Link", strategic_plan_link or ""),
        ("Total Public Meeting Time", f"{total_min} minutes"),
        ("Goal-Focused Time", f"{gf_min} minutes ({gf_pct}% — target: ≥50%)"),
    ]
    for label, value in rows_data:
        _add_info_row(t1, label, value)

    # ── Section 2: Agenda Item Evaluation ─────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    _section_heading(doc, "Section 2: Agenda Item Evaluation")
    t2 = doc.add_table(rows=1, cols=4)
    t2.style = "Table Grid"
    t2.columns[0].width = Inches(0.4)
    t2.columns[1].width = Inches(2.9)
    t2.columns[2].width = Inches(0.5)
    t2.columns[3].width = Inches(3.2)
    _header_row(t2, ["#", "Agenda Item", "Min", "Category / Notes"])

    item_map = {i["item_number"]: i for i in agenda_items}
    cls_map = {c["item_number"]: c for c in classified_items}

    for num in sorted(item_map.keys()):
        item = item_map[num]
        cls = cls_map.get(num, {})
        row = t2.add_row()
        cells = row.cells
        cells[0].paragraphs[0].add_run(str(num)).font.size = Pt(8)
        cells[1].paragraphs[0].add_run(item.get("title", "")).font.size = Pt(8)
        cells[2].paragraphs[0].add_run(str(cls.get("minutes", 0))).font.size = Pt(8)
        note_text = cls.get("category", "Other")
        if cls.get("notes"):
            note_text += f" — {cls['notes']}"
        cells[3].paragraphs[0].add_run(note_text).font.size = Pt(8)
        if cls.get("goal_focused"):
            for cell in cells:
                _set_cell_bg(cell, ESB_LIGHT)

    # Total row
    total_row = t2.add_row()
    for cell in total_row.cells:
        _set_cell_bg(cell, ESB_DARK)
    total_row.cells[0].paragraphs[0].add_run("").font.size = Pt(8)
    tr2 = total_row.cells[1].paragraphs[0].add_run("TOTAL PUBLIC MEETING")
    tr2.bold = True
    tr2.font.color.rgb = ESB_WHITE
    tr2.font.size = Pt(8)
    total_row.cells[2].paragraphs[0].add_run(str(total_min)).font.color.rgb = ESB_WHITE
    total_row.cells[2].paragraphs[0].runs[0].font.size = Pt(8)
    gf_run = total_row.cells[3].paragraphs[0].add_run(f"Goal-Focused: {gf_min} min ({gf_pct}%)")
    gf_run.bold = True
    gf_run.font.color.rgb = ESB_WHITE
    gf_run.font.size = Pt(8)

    # ── Section 3: Category Summary ───────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    _section_heading(doc, "Section 3: Time Use by Category")

    CATEGORY_ORDER = [
        ("Goal Setting", True),
        ("Goal Monitoring", True),
        ("Guardrail Setting", False),
        ("Guardrail Monitoring", False),
        ("Data Evaluation (Student Data)", False),
        ("Data Evaluation (System Data)", False),
        ("Superintendent Evaluation", False),
        ("Community Listening (Goals)", False),
        ("Community Listening (Guardrails)", False),
        ("Policy Review", False),
        ("Budget Review", False),
        ("Board Self Evaluation", False),
        ("Board Training", False),
        ("Voting / Non-Goal Actions", False),
        ("Closed Session", False),
        ("Other", False),
    ]

    t3 = doc.add_table(rows=1, cols=4)
    t3.style = "Table Grid"
    t3.columns[0].width = Inches(2.9)
    t3.columns[1].width = Inches(0.7)
    t3.columns[2].width = Inches(0.8)
    t3.columns[3].width = Inches(2.6)
    _header_row(t3, ["Category", "Minutes", "% of Total", "Notes"])

    by_cat = category_totals["by_category"]
    for cat_name, goal_focused in CATEGORY_ORDER:
        mins = by_cat.get(cat_name, 0)
        pct = round(mins / total_min * 100, 1) if total_min else 0
        row = t3.add_row()
        cells = row.cells
        label_run = cells[0].paragraphs[0].add_run(cat_name)
        label_run.font.size = Pt(8)
        if goal_focused:
            label_run.bold = True
        cells[1].paragraphs[0].add_run(str(mins) if mins else "—").font.size = Pt(8)
        cells[2].paragraphs[0].add_run(f"{pct}%" if mins else "—").font.size = Pt(8)
        note = "✓ Counts toward 50% target" if goal_focused and mins else ""
        cells[3].paragraphs[0].add_run(note).font.size = Pt(8)
        if goal_focused and mins:
            for cell in cells:
                _set_cell_bg(cell, ESB_LIGHT)

    tot_row = t3.add_row()
    for cell in tot_row.cells:
        _set_cell_bg(cell, ESB_DARK)
    tr3 = tot_row.cells[0].paragraphs[0].add_run("TOTAL")
    tr3.bold = True
    tr3.font.color.rgb = ESB_WHITE
    tr3.font.size = Pt(8)
    tot_row.cells[1].paragraphs[0].add_run(str(total_min)).font.color.rgb = ESB_WHITE
    tot_row.cells[1].paragraphs[0].runs[0].font.size = Pt(8)
    tot_row.cells[2].paragraphs[0].add_run("100%").font.color.rgb = ESB_WHITE
    tot_row.cells[2].paragraphs[0].runs[0].font.size = Pt(8)
    gf_s = tot_row.cells[3].paragraphs[0].add_run(f"Goal-Focused: {gf_min} min ({gf_pct}%)")
    gf_s.bold = True
    gf_s.font.color.rgb = ESB_WHITE
    gf_s.font.size = Pt(8)

    # ── Section 4: Coaching Reflections ───────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    _section_heading(doc, "Section 4: Coaching Reflections")

    for section_key, section_label in [
        ("celebrates", "Practitioner Celebrates"),
        ("recommends", "Practitioner Recommends"),
    ]:
        label_para = doc.add_paragraph()
        lr = label_para.add_run(section_label)
        lr.bold = True
        lr.font.color.rgb = ESB_BLUE
        lr.font.size = Pt(10)
        label_para.paragraph_format.space_before = Pt(6)
        label_para.paragraph_format.space_after = Pt(2)

        for entry in coaching.get(section_key, []):
            h_para = doc.add_paragraph()
            h_para.paragraph_format.space_before = Pt(3)
            h_para.paragraph_format.space_after = Pt(1)
            hr = h_para.add_run(entry["heading"])
            hr.bold = True
            hr.font.color.rgb = ESB_DARK
            hr.font.size = Pt(9)

            b_para = doc.add_paragraph()
            b_para.paragraph_format.space_before = Pt(0)
            b_para.paragraph_format.space_after = Pt(3)
            b_para.paragraph_format.left_indent = Inches(0.2)
            br = b_para.add_run(entry["body"])
            br.font.size = Pt(8)

    doc.save(output_path)


def generate_multi_meeting_docx(
    output_path: str,
    district: str,
    span_label: str,
    scored_meetings: List[Dict],
    aggregate: Dict,
    coaching: Dict[str, List[Dict]],
):
    """Generate a multi-meeting governance review document (schoolboardreview methodology)."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.3)
        _setup_header(section)
        _setup_footer(section)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(2)
    tr = title.add_run(f"ESB Governance Review — {span_label}")
    tr.bold = True
    tr.font.size = Pt(14)
    tr.font.color.rgb = ESB_DARK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(6)
    sr = subtitle.add_run(district)
    sr.font.size = Pt(11)
    sr.font.color.rgb = ESB_BLUE

    # Section 1: Review Information
    _section_heading(doc, "Section 1: Review Information")
    t1 = doc.add_table(rows=1, cols=2)
    t1.style = "Table Grid"
    t1.columns[0].width = Inches(2.2)
    t1.columns[1].width = Inches(4.8)
    _header_row(t1, ["Field", "Value"])

    dates = sorted(m.get("meeting_date", "?") for m in scored_meetings if m.get("meeting_date"))
    date_range = f"{dates[0]} — {dates[-1]}" if len(dates) > 1 else (dates[0] if dates else "—")
    composite_avg = aggregate.get("composite_avg", 0)

    for label, value in [
        ("School System", district),
        ("Review Span", span_label),
        ("Meetings Analyzed", str(len(scored_meetings))),
        ("Date Range", date_range),
        ("Composite Score (avg)", f"{composite_avg:.1f} / 100"),
    ]:
        _add_info_row(t1, label, value)

    # Section 2: Governance Score Breakdown
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    _section_heading(doc, "Section 2: Governance Score Breakdown (avg across all meetings)")

    t2 = doc.add_table(rows=1, cols=4)
    t2.style = "Table Grid"
    t2.columns[0].width = Inches(2.0)
    t2.columns[1].width = Inches(1.2)
    t2.columns[2].width = Inches(0.8)
    t2.columns[3].width = Inches(3.0)
    _header_row(t2, ["Criterion", "Avg Score", "Trend", "Description"])

    for c in GOVERNANCE_CRITERIA:
        key = c["key"]
        info = aggregate.get(key, {"avg": 0, "trend": "—"})
        row = t2.add_row()
        cells = row.cells
        cr = cells[0].paragraphs[0].add_run(c["name"])
        cr.bold = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = ESB_DARK
        cells[1].paragraphs[0].add_run(f"{info['avg']:.1f} / 20").font.size = Pt(9)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        trend_run = cells[2].paragraphs[0].add_run(info["trend"])
        trend_run.font.size = Pt(8)
        trend_run.font.color.rgb = (
            RGBColor(0x05, 0x96, 0x69) if info["trend"] == "Improving"
            else RGBColor(0xDC, 0x26, 0x26) if info["trend"] == "Declining"
            else ESB_GRAY
        )
        cells[3].paragraphs[0].add_run(c["desc"]).font.size = Pt(8)

    # Total row
    tot = t2.add_row()
    for cell in tot.cells:
        _set_cell_bg(cell, ESB_DARK)
    tr2 = tot.cells[0].paragraphs[0].add_run("COMPOSITE AVERAGE")
    tr2.bold = True
    tr2.font.color.rgb = ESB_WHITE
    tr2.font.size = Pt(9)
    cr2 = tot.cells[1].paragraphs[0].add_run(f"{composite_avg:.1f} / 100")
    cr2.bold = True
    cr2.font.color.rgb = ESB_WHITE
    cr2.font.size = Pt(9)
    tot.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Section 3: Meeting-by-Meeting Scores
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    _section_heading(doc, "Section 3: Meeting-by-Meeting Score Detail")

    col_headers = ["Date"] + [c["name"] for c in GOVERNANCE_CRITERIA] + ["Composite"]
    col_widths = [Inches(0.9)] + [Inches(0.9)] * 5 + [Inches(0.9)]

    t3 = doc.add_table(rows=1, cols=len(col_headers))
    t3.style = "Table Grid"
    for i, w in enumerate(col_widths):
        t3.columns[i].width = w
    _header_row(t3, col_headers)

    for m in sorted(scored_meetings, key=lambda x: x.get("meeting_date", "")):
        row = t3.add_row()
        cells = row.cells
        cells[0].paragraphs[0].add_run(m.get("meeting_date", "?")).font.size = Pt(8)
        for i, c in enumerate(GOVERNANCE_CRITERIA, 1):
            score_val = m.get(c["key"], {}).get("score", "—")
            cells[i].paragraphs[0].add_run(
                f"{score_val}/20" if isinstance(score_val, int) else str(score_val)
            ).font.size = Pt(8)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        comp = m.get("composite", "—")
        comp_run = cells[6].paragraphs[0].add_run(
            f"{comp}/100" if isinstance(comp, int) else str(comp)
        )
        comp_run.bold = True
        comp_run.font.size = Pt(8)
        cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Section 4: Coaching Reflections
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    _section_heading(doc, "Section 4: Coaching Reflections")

    for section_key, section_label in [
        ("celebrates", "Practitioner Celebrates"),
        ("recommends", "Practitioner Recommends"),
    ]:
        label_para = doc.add_paragraph()
        lr = label_para.add_run(section_label)
        lr.bold = True
        lr.font.color.rgb = ESB_BLUE
        lr.font.size = Pt(10)
        label_para.paragraph_format.space_before = Pt(6)
        label_para.paragraph_format.space_after = Pt(2)

        for entry in coaching.get(section_key, []):
            h_para = doc.add_paragraph()
            h_para.paragraph_format.space_before = Pt(3)
            h_para.paragraph_format.space_after = Pt(1)
            hr = h_para.add_run(entry["heading"])
            hr.bold = True
            hr.font.color.rgb = ESB_DARK
            hr.font.size = Pt(9)

            b_para = doc.add_paragraph()
            b_para.paragraph_format.space_before = Pt(0)
            b_para.paragraph_format.space_after = Pt(3)
            b_para.paragraph_format.left_indent = Inches(0.2)
            br = b_para.add_run(entry["body"])
            br.font.size = Pt(8)

    doc.save(output_path)




def _shaded_box(doc, heading, body_parts, bg, heading_color):
    """Render a highlight box: shaded 1-column table with bold heading + body paragraphs/quotes."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, bg)
    if heading:
        hp = cell.paragraphs[0]
        hr = hp.add_run(heading)
        hr.bold = True
        hr.font.size = Pt(9)
        hr.font.color.rgb = heading_color
    for chunk in body_parts:
        if chunk["type"] == "quote":
            qp = cell.add_paragraph()
            qp.paragraph_format.left_indent = Inches(0.18)
            qp.paragraph_format.space_before = Pt(4)
            qr = qp.add_run('"' + chunk["text"] + '"')
            qr.italic = True
            qr.font.size = Pt(8.5)
            qr.font.color.rgb = ESB_DARK
            sp = cell.add_paragraph()
            sp.paragraph_format.left_indent = Inches(0.18)
            sr = sp.add_run("— " + chunk["source"])
            sr.font.size = Pt(7.5)
            sr.font.color.rgb = ESB_GRAY
            sr.bold = True
        else:
            content = chunk.get("content", "").strip()
            if content:
                bp = cell.add_paragraph()
                bp.paragraph_format.space_before = Pt(3)
                brr = bp.add_run(content)
                brr.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


